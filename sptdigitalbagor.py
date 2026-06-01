import streamlit as st
from streamlit_drawable_canvas import st_canvas
from google.oauth2 import service_account
from googleapiclient.discovery import build
import datetime
from io import BytesIO
from PIL import Image
from docx import Document
from docx.shared import Mm
import os
import base64

INDO_MONTHS = ["Januari", "Februari", "Maret", "April", "Mei", "Juni", "Juli", "Agustus", "September", "Oktober", "November", "Desember"]

def format_indo_date(date_obj):
    day = date_obj.day
    month = INDO_MONTHS[date_obj.month - 1]
    year = date_obj.year
    return f"{day} {month} {year}"

# --- 1. SETTING HALAMAN ---
st.set_page_config(page_title="Kirim Surat Tugas", layout="wide", page_icon="📝")

# --- 2. KONEKSI GOOGLE SHEETS ---
@st.cache_resource
def get_sheets_service():
    try:
        if "gcp_service_account" in st.secrets:
            cred_info = st.secrets["gcp_service_account"]
            creds = service_account.Credentials.from_service_account_info(cred_info)
            return build('sheets', 'v4', credentials=creds)
        return None
    except:
        return None

sheets_service = get_sheets_service()
SPREADSHEET_ID = "1hA68rgMDtbX9ySdOI5TF5CUypzO5vJKHHIPAVjTk798"

# --- 2B. FUNGSI DATABASE (KEGIATAN & SUBMISI) ---
def get_kegiatan_list():
    if not sheets_service:
        return [
            {"nama": "SPT Rekon TPP dan SIMONA", "integrasi": "SITPP", "status": "Aktif", "deadline": "6 Februari 2026"},
            {"nama": "Lainnya", "integrasi": "None", "status": "Aktif", "deadline": "Tanpa Batas"}
        ]
    try:
        meta = sheets_service.spreadsheets().get(spreadsheetId=SPREADSHEET_ID).execute()
        sheets = [s['properties']['title'] for s in meta['sheets']]
        
        if "Config_Kegiatan" not in sheets:
            # Buat sheet baru
            body = {
                'requests': [{
                    'addSheet': {
                        'properties': {
                            'title': 'Config_Kegiatan'
                        }
                    }
                }]
            }
            sheets_service.spreadsheets().batchUpdate(spreadsheetId=SPREADSHEET_ID, body=body).execute()
            
            # Tulis data default
            default_rows = [
                ["Nama Kegiatan", "Integrasi", "Status", "Batas Tanggal"],
                ["SPT Rekon TPP dan SIMONA", "SITPP", "Aktif", "6 Februari 2026"],
                ["Lainnya", "None", "Aktif", "Tanpa Batas"]
            ]
            sheets_service.spreadsheets().values().update(
                spreadsheetId=SPREADSHEET_ID,
                range="Config_Kegiatan!A1",
                valueInputOption="USER_ENTERED",
                body={'values': default_rows}
            ).execute()
            return [
                {"nama": "SPT Rekon TPP dan SIMONA", "integrasi": "SITPP", "status": "Aktif", "deadline": "6 Februari 2026"},
                {"nama": "Lainnya", "integrasi": "None", "status": "Aktif", "deadline": "Tanpa Batas"}
            ]
        
        result = sheets_service.spreadsheets().values().get(
            spreadsheetId=SPREADSHEET_ID,
            range="Config_Kegiatan!A2:D100"
        ).execute()
        rows = result.get('values', [])
        if not rows:
            return []
        
        kegiatan = []
        for r in rows:
            if len(r) >= 1:
                kegiatan.append({
                    "nama": r[0],
                    "integrasi": r[1] if len(r) > 1 else "None",
                    "status": r[2] if len(r) > 2 else "Aktif",
                    "deadline": r[3] if len(r) > 3 else "Tanpa Batas"
                })
        return kegiatan
    except Exception as e:
        return [
            {"nama": "SPT Rekon TPP dan SIMONA", "integrasi": "SITPP", "status": "Aktif", "deadline": "6 Februari 2026"},
            {"nama": "Lainnya", "integrasi": "None", "status": "Aktif", "deadline": "Tanpa Batas"}
        ]

def save_kegiatan_list(kegiatan_list):
    if not sheets_service:
        return False
    try:
        rows = [["Nama Kegiatan", "Integrasi", "Status", "Batas Tanggal"]]
        for k in kegiatan_list:
            rows.append([k["nama"], k["integrasi"], k["status"], k.get("deadline", "Tanpa Batas")])
        
        # Bersihkan data lama
        sheets_service.spreadsheets().values().clear(
            spreadsheetId=SPREADSHEET_ID,
            range="Config_Kegiatan!A1:D100"
        ).execute()
        
        # Tulis data baru
        sheets_service.spreadsheets().values().update(
            spreadsheetId=SPREADSHEET_ID,
            range="Config_Kegiatan!A1",
            valueInputOption="USER_ENTERED",
            body={'values': rows}
        ).execute()
        return True
    except Exception as e:
        st.error(f"Gagal menyimpan konfigurasi kegiatan: {e}")
        return False


def get_submissions_data():
    if not sheets_service:
        return []
    try:
        result = sheets_service.spreadsheets().values().get(
            spreadsheetId=SPREADSHEET_ID,
            range="Sheet1!A1:M1000"
        ).execute()
        return result.get('values', [])
    except Exception as e:
        return []


# --- 3. FUNGSI KONVERSI TTD KE BASE64 (Untuk Spreadsheet) ---
def get_base64_signature(signature_img):
    try:
        if signature_img is not None and signature_img.any():
            img_rgba = Image.fromarray(signature_img.astype('uint8'), 'RGBA')
            # Tambahkan background putih agar tidak transparan di Base64
            white_bg = Image.new("RGBA", img_rgba.size, (255, 255, 255, 255))
            final_img = Image.alpha_composite(white_bg, img_rgba).convert("RGB")
            
            buffered = BytesIO()
            final_img.save(buffered, format="PNG")
            return base64.b64encode(buffered.getvalue()).decode()
        return "Tidak ada TTD"
    except:
        return "Error TTD"

# --- 4. FUNGSI GENERATE DOCX (Untuk Download) ---
def create_docx_final(data, signature_img):
    template_name = "template spt simona.docx" 
    if not os.path.exists(template_name):
        st.error("File template tidak ditemukan!")
        return None
    try:
        doc = Document(template_name)
        replacements = {
            '{{unitkerja}}': str(data['unit_kerja']),
            '{{nama_admin}}': str(data['nama']),
            '{{pangkat_admin}}': str(data['pangkat']),
            '{{NIP_admin}}': str(data['nip']),
            '{{Jabatanadmin}}': str(data['jabatan']),
            '{{no_hpadmin}}': str(data['no_hp']),
            '{{email_admin}}': str(data['email']),
            '{{JABATAN_ATASAN}}': str(data['j_atasan']),
            '{{NAMA_ATASAN}}': str(data['n_atasan']),
            '{{NIP_ATASAN}}': str(data['nip_atasan']),
            '{{PANGKAT_GOL_ATASAN}}': str(data['p_atasan']),
            '{{perihal}}': str(data['perihal']),
            '{{TTL}}': datetime.datetime.now().strftime('%d %B %Y')
        }

        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    for run in paragraph.runs:
                        if key in run.text:
                            run.text = run.text.replace(key, value)
            
            if '{{ttd}}' in paragraph.text:
                for run in paragraph.runs:
                    if '{{ttd}}' in run.text:
                        run.text = run.text.replace('{{ttd}}', "")
                if signature_img is not None:
                    img_rgba = Image.fromarray(signature_img.astype('uint8'), 'RGBA')
                    white_bg = Image.new("RGBA", img_rgba.size, (255, 255, 255, 255))
                    final_img = Image.alpha_composite(white_bg, img_rgba).convert("RGB")
                    img_io = BytesIO()
                    final_img.save(img_io, format='PNG')
                    img_io.seek(0)
                    new_run = paragraph.add_run()
                    new_run.add_picture(img_io, width=Mm(45))

        target_stream = BytesIO()
        doc.save(target_stream)
        target_stream.seek(0)
        return target_stream
    except Exception as e:
        st.error(f"Gagal memproses dokumen: {e}")
        return None

# --- 5. TAMPILAN OPERATOR FORM ---
def show_operator_form():
    # TAMBAHAN: FUNGSI MODAL YANG LEBIH STABIL
    @st.dialog("PENGUMUMAN PENTING")
    def tampilkan_pengumuman():
        st.warning("⚠️ **Batas Pengiriman SPT:**")
        
        keg_list = get_kegiatan_list()
        active_kegs = [k for k in keg_list if k.get("status") == "Aktif"]
        
        if active_kegs:
            for k in active_kegs:
                dl = k.get("deadline", "Tanpa Batas")
                st.markdown(f"- **{k['nama']}**: batas tanggal **{dl}**")
        else:
            st.markdown("- Tidak ada kegiatan aktif saat ini.")
            
        st.info("💡 **Informasi:**\n\nTidak perlu menyerahkan SPT Fisik ke Bagian Organisasi.")
        if st.button("Saya Mengerti", type="primary", width="stretch"):
            st.session_state["sudah_baca_info"] = True
            st.rerun()

    # Logika agar hanya muncul satu kali saat aplikasi pertama kali dimuat
    if "sudah_baca_info" not in st.session_state:
        tampilkan_pengumuman()
    st.markdown("<h2 style='text-align: center;'>📝 Kirim Surat Tugas</h2>", unsafe_allow_html=True)
    st.write("---")

    # SEKSI I: PERIHAL & OPD
    st.subheader("I. Perihal & Unit Kerja")
    
    kegiatan_list = get_kegiatan_list()
    # Hanya cari kegiatan yang statusnya Aktif (selain "Lainnya" yang mungkin dimasukkan admin)
    nama_kegiatan_active = [k["nama"] for k in kegiatan_list if k["status"] == "Aktif" and k["nama"] != "Lainnya"]
    
    is_disabled = False
    if not nama_kegiatan_active:
        nama_kegiatan_active = ["Tidak ada kegiatan aktif saat ini"]
        is_disabled = True
        st.warning("⚠️ **Pengiriman SPT Dinonaktifkan**: Tidak ada kegiatan yang di-input atau aktif di sistem saat ini.")
    else:
        if "Lainnya" not in nama_kegiatan_active:
            nama_kegiatan_active.append("Lainnya")
        
    opsi_perihal = st.selectbox("Pilih Perihal / Kegiatan:", nama_kegiatan_active, disabled=is_disabled)
    
    if is_disabled:
        perihal_final = ""
    else:
        perihal_final = st.text_input("Ketik Perihal Manual (Jika Lainnya):") if opsi_perihal == "Lainnya" else opsi_perihal
    
    selected_kegiatan = next((k for k in kegiatan_list if k["nama"] == opsi_perihal), None)
    integrasi_val = selected_kegiatan["integrasi"] if selected_kegiatan else "None"
    if opsi_perihal == "Lainnya" or is_disabled:
        integrasi_val = "None"
        
    current_year_int = datetime.datetime.now().year
    target_tahun = st.selectbox("Tahun Kegiatan:", [str(y) for y in range(current_year_int - 1, current_year_int + 5)], index=1, disabled=is_disabled)

    list_opd = [
        "Bagian Tata Pemerintahan", "Bagian Kesejahteraan Rakyat", "Bagian Hukum", "Bagian Kerjasama", "Bagian Perekonomian", "Bagian Pembangunan dan Sumber Daya Alam", "Bagian Pengadaan Barang dan Jasa", "Bagian Umum", "Bagian Organisasi", "Bagian Protokol dan Komunikasi Pimpinan", "Bagian Perencanaan dan Keuangan", "Sekretariat DPRD", "Inspektorat Daerah",
        "Dinas Pendidikan dan Kebudayaan", "Dinas Kesehatan", "Dinas Pekerjaan Umum dan Penataan Ruang",
        "Dinas Perumahan dan Kawasan Permukiman", "Satuan Polisi Pamong Praja", "Dinas Pemadam Kebakaran dan Penyelamatan",
        "Dinas Sosial, Pemberdayaan Perempuan dan Perlindungan Anak", "Dinas Lingkungan Hidup",
        "Dinas Kependudukan dan Pencatatan Sipil", "Dinas Pemberdayaan Masyarakat dan Desa",
        "Dinas Perhubungan", "Dinas Komunikasi dan Informatika", "Dinas Koperasi, Perindustrian dan Perdagangan",
        "Dinas Penanaman Modal dan Pelayanan Terpadu Satu Pintu", "Dinas Pariwisata, Pemuda dan Olahraga",
        "Dinas Perpustakaan dan Arsip Daerah", "Dinas Perikanan", "Dinas Ketahanan Pangan",
        "Dinas Tanaman Pangan dan Hortikultura", "Dinas Perkebunan dan Peternakan",
        "Dinas Tenaga Kerja dan Transmigrasi", "Dinas Pengendalian Penduduk dan Keluarga Berencana",
        "Badan Perencanaan Pembanguan Dan Riset Inovasi Daerah", "Badan Pengelola Keuangan dan Aset Daerah", 
        "Badan Pengelola Pajak dan Retribusi Daerah", "Badan Kepegawaian dan Pengembangan Sumber Daya Manusia ", 
        "Badan Penanggulangan Bencana Daerah", "Kesbangpol",
        "RSUD Ahmad Ripin", "RSUD Sungai Gelam", "RSUD Sungai Bahar",
        "Kecamatan Sekernan", "Kecamatan Jaluko", "Kecamatan Maro Sebo", "Kecamatan Kumpeh",
        "Kecamatan Kumpeh Ulu", "Kecamatan Mestong", "Kecamatan Sungai Gelam", "Kecamatan Sungai Bahar",
        "Kecamatan Bahar Utara", "Kecamatan Bahar Selatan", "Kecamatan Taman Rajo", "Puskesmas Penyengat Olak", "Puskesmas Kemingking Dalam", "Puskesmas Puding", "Puskesmas Sungai Bahar IV", "Puskesmas Simpang Sungai Duren", "Puskesmas Tempino", "Puskesmas Tanjung", "Puskesmas Bahar VII", "Puskesmas Tantan", "Puskesmas Kasang Pudak", "Puskesmas Sengeti", "Puskesmas Pir II Bajubang", "Puskesmas Pondok Meja", "Puskesmas Markanding", "Puskesmas Tangkit", "Puskesmas Talang Bukit", "Puskesmas Sekernan Ilir", "Puskesmas Jambi Kecil", "Puskesmas Muara Kumpeh", "Puskesmas Sungai Bahar I", "Puskesmas Kebon IX", "Puskesmas Suko Awin", "Puskesmas Petaling Jaya"
    ]
    opsi_opd = st.selectbox("Pilih Unit Kerja / OPD:", [""] + sorted(list_opd) + ["Lainnya"], disabled=is_disabled)
    if is_disabled:
        unit_kerja_final = ""
    else:
        unit_kerja_final = st.text_input("Ketik Nama OPD (Jika Lainnya):") if opsi_opd == "Lainnya" else opsi_opd

    st.write("---")

    # SEKSI II: DATA ADMIN
    st.subheader("II. Data Admin")
    status_pegawai = st.radio("Status Pegawai:", ["PNS", "PPPK"], horizontal=True, disabled=is_disabled)

    c1, c2 = st.columns(2)
    with c1:
        nama_admin = st.text_input("Nama Lengkap", disabled=is_disabled)
        nip_admin = st.text_input(f"NIP / NI {status_pegawai}", max_chars=18, placeholder="18 Digit Angka", disabled=is_disabled)
        no_hp = st.text_input("Nomor WhatsApp", disabled=is_disabled)
    with c2:
        pangkat_admin = st.text_input("Pangkat / Golongan", disabled=is_disabled)
        jabatan_admin = st.text_input("Jabatan", disabled=is_disabled)
        email = st.text_input("Email", placeholder="harus @gmail.com", disabled=is_disabled)

    st.write("---")

    # SEKSI III: DATA ATASAN
    st.subheader("III. Data Atasan")
    n_atasan = st.text_input("Nama Lengkap Atasan", disabled=is_disabled)
    j_atasan = st.text_input("Jabatan Atasan (Contoh: Kepala Bagian Organisasi)", disabled=is_disabled)

    c3, c4 = st.columns(2)
    with c3:
        p_atasan = st.text_input("Pangkat / Golongan Atasan", disabled=is_disabled)
    with c4:
        nip_atasan = st.text_input("NIP Atasan", max_chars=18, disabled=is_disabled)
        st.info(f"Tanggal Surat: {datetime.datetime.now().strftime('%d %B %Y')}")

    st.write("---")

    # SEKSI IV: TANDA TANGAN
    st.subheader("IV. Tanda Tangan Atasan")

    canvas_result = st_canvas(
        stroke_width=3, 
        stroke_color="#000000", 
        background_color="#ffffff",
        height=150, 
        width=350, 
        drawing_mode="freedraw", 
        key="canvas_final",
        display_toolbar=True
    )

    st.markdown("""
        <p style='color: #ff4b4b; font-size: 0.85rem; font-weight: bold; margin-top: -10px;'>
            ⚠️ Pastikan kolom di atas ditandatangani oleh Atasan yang bersangkutan! (Contoh: kepala dinas/badan)
        </p>
        """, unsafe_allow_html=True)

    st.write("")
    if st.button("KIRIM DATA", type="primary", width="stretch", disabled=is_disabled):
        val_nip = nip_admin.isdigit() and len(nip_admin) == 18 and nip_atasan.isdigit() and len(nip_atasan) == 18
        val_email = email.lower().endswith("@gmail.com")

        if not val_nip:
            st.error("❌ NIP Admin dan Atasan harus 18 digit angka!")
        elif not val_email:
            st.error("❌ Email wajib menggunakan domain @gmail.com!")
        elif not nama_admin or not unit_kerja_final or not n_atasan:
            st.warning("⚠️ Mohon lengkapi semua field yang tersedia!")
        else:
            with st.spinner('Memproses data...'):
                ttd_b64 = get_base64_signature(canvas_result.image_data)
                
            if sheets_service:
                try:
                    now = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                    
                    row = [[
                        now,                    # Kolom A: Waktu
                        perihal_final,          # Kolom B: Perihal
                        unit_kerja_final,       # Kolom C: Unit Kerja
                        f"({status_pegawai}) {nama_admin}", # Kolom D: Nama Admin
                        f"'{nip_admin}",        # Kolom E: NIP Admin
                        email,                  # Kolom F: Email
                        n_atasan,               # Kolom G: Nama Atasan
                        j_atasan,               # Kolom H: Jabatan Atasan
                        p_atasan,               # Kolom I: Pangkat Atasan
                        f"'{nip_atasan}",       # Kolom J: NIP Atasan
                        ttd_b64,                # Kolom K: Base64 TTD
                        integrasi_val,          # Kolom L: Integrasi
                        str(target_tahun)       # Kolom M: Tahun
                    ]]
                    
                    sheets_service.spreadsheets().values().append(
                        spreadsheetId=SPREADSHEET_ID, 
                        range="Sheet1!A1",
                        valueInputOption="USER_ENTERED", 
                        body={'values': row}
                    ).execute()
                    
                except Exception as e:
                    st.error(f"Gagal kirim ke Sheets: {e}")

                data_spt = {
                    'unit_kerja': unit_kerja_final, 'nama': nama_admin, 'nip': nip_admin,
                    'pangkat': pangkat_admin, 'jabatan': jabatan_admin, 'no_hp': no_hp,
                    'email': email, 'j_atasan': j_atasan, 'n_atasan': n_atasan,
                    'nip_atasan': nip_atasan, 'p_atasan': p_atasan, 'perihal': perihal_final
                }
                docx_file = create_docx_final(data_spt, canvas_result.image_data)
                
                if docx_file:
                    st.success("✅ Data berhasil masuk, Terima Kasih")
                    st.download_button("📥 Download SPT Sekarang", docx_file, f"SPT_{nama_admin.replace(' ','_')}.docx", width="stretch")

# --- 5B. TAMPILAN ADMIN DASHBOARD ---
def show_admin_page():
    st.markdown("<h2 style='text-align: center;'>🔒 Panel Admin SPT Digital</h2>", unsafe_allow_html=True)
    st.write("---")
    
    correct_password = st.secrets.get("admin_password", "adminbagor123")
    
    if "admin_authenticated" not in st.session_state:
        st.session_state["admin_authenticated"] = False
        
    if not st.session_state["admin_authenticated"]:
        passwd = st.text_input("Masukkan Password Admin:", type="password")
        if st.button("LOG IN", type="primary", width="stretch"):
            if passwd == correct_password:
                st.session_state["admin_authenticated"] = True
                st.success("Login Berhasil!")
                st.rerun()
            else:
                st.error("Password salah!")
        return

    if st.sidebar.button("Log Out Admin"):
        st.session_state["admin_authenticated"] = False
        st.rerun()
        
    tab1, tab2 = st.tabs(["📊 Dashboard & Data Submisi", "⚙️ Pengaturan Kegiatan (Perihal)"])
    
    with tab1:
        with st.spinner("Memuat data dari Google Sheets..."):
            rows = get_submissions_data()
        
        if not rows or len(rows) <= 1:
            st.info("Belum ada data submisi SPT.")
        else:
            data_rows = rows[1:]
            
            import pandas as pd
            
            # Tentukan jumlah kolom maksimum dari data
            max_cols = max(len(rows[0]), max(len(r) for r in data_rows) if data_rows else 0)
            
            # Definisikan nama kolom standar berdasarkan posisi index kolom di Google Sheet
            standard_headers = [
                "Waktu", "Perihal", "Unit Kerja", "Nama Admin", "NIP Admin", 
                "Email", "Nama Atasan", "Jabatan Atasan", "Pangkat Gol Atasan", 
                "NIP Atasan", "TTD", "Integrasi", "Tahun"
            ]
            
            final_headers = []
            for idx in range(max_cols):
                if idx < len(standard_headers):
                    final_headers.append(standard_headers[idx])
                else:
                    final_headers.append(f"Kolom_{idx+1}")
            
            padded_data = []
            for r in data_rows:
                padded_r = r + [""] * (max_cols - len(r))
                padded_data.append(padded_r)
                
            df = pd.DataFrame(padded_data, columns=final_headers)
            
            # Pastikan kolom Integrasi dan Tahun selalu ada (terutama untuk data lama)
            if "Integrasi" not in df.columns:
                df["Integrasi"] = "None"
            if "Tahun" not in df.columns:
                def get_year_from_time(t):
                    try:
                        return str(t).split("-")[0].strip()
                    except:
                        return str(datetime.datetime.now().year)
                df["Tahun"] = df["Waktu"].apply(get_year_from_time)
            
            df["Tahun"] = df["Tahun"].apply(lambda x: str(x).strip() if x else str(datetime.datetime.now().year))
            df["Integrasi"] = df["Integrasi"].apply(lambda x: str(x).strip() if x else "None")
            
            available_years = sorted(list(df["Tahun"].unique()))
            if not available_years:
                available_years = [str(datetime.datetime.now().year)]
                
            selected_year = st.selectbox("Pilih Tahun Kegiatan:", available_years, index=len(available_years)-1)
            
            df_filtered = df[df["Tahun"] == selected_year]
            
            # --- Tambahan Filter Nama Kegiatan ---
            kegiatan_db_names = [k["nama"] for k in get_kegiatan_list()]
            kegiatan_data_names = list(df_filtered["Perihal"].unique()) if not df_filtered.empty else []
            available_kegiatan = ["Semua Kegiatan"] + sorted(list(set(kegiatan_db_names + kegiatan_data_names)))
            
            selected_kegiatan = st.selectbox("Pilih Nama Kegiatan:", available_kegiatan, index=0)
            if selected_kegiatan != "Semua Kegiatan":
                df_filtered = df_filtered[df_filtered["Perihal"] == selected_kegiatan]
            
            # --- Perhitungan compliance & stats ---
            masterUnit = {
                "Dinas/Badan": [
                    "Sekretariat Dewan Perwakilan Rakyat Daerah", "Inspektorat Daerah", "Dinas Pendidikan dan Kebudayaan", 
                    "Dinas Pariwisata, Kepemudaan dan Olahraga", "Dinas Kesehatan", "Dinas Sosial, Pemberdayaan Perempuan dan Perlindungan Anak", 
                    "Dinas Pengendalian Penduduk dan Keluarga Berencana", "Dinas Kependudukan dan Pencatatan Sipil", 
                    "Dinas Pemberdayaan Masyarakat dan Desa", "Satuan Polisi Pamong Praja", "Dinas Penanaman Modal dan Pelayanan Terpadu Satu Pintu", 
                    "Dinas Koperasi, Usaha Kecil Menengah, Perindustrian dan Perdagangan", "Dinas Tenaga Kerja dan Transmigrasi", 
                    "Dinas Komunikasi dan Informatika", "Dinas Perumahan Dan Kawasan Permukiman", "Dinas Pekerjaan Umum dan Penataan Ruang", 
                    "Dinas Perhubungan", "Dinas Lingkungan Hidup", "Dinas Tanaman Pangan dan Hortikultura", "Dinas Ketahanan Pangan", 
                    "Dinas Perkebunan dan Peternakan", "Dinas Perikanan", "Dinas Perpustakaan dan Arsip Daerah", 
                    "Badan Perencanaan Pembanguan Dan Riset Inovasi Daerah", "Badan Kepegawaian dan Pengembangan Sumber Daya Manusia", 
                    "Badan Pengelola Keuangan dan Aset Daerah", "Badan Pengelola Pajak dan Retribusi Daerah", 
                    "Badan Penanggulangan Bencana Daerah", "Dinas Pemadam Kebakaran dan Penyelamatan", "Kesbangpol"
                ],
                "Bagian (Setda)": [
                    "Bagian Tata Pemerintahan", "Bagian Kesejahteraan Rakyat", "Bagian Hukum", "Bagian Kerjasama", 
                    "Bagian Perekonomian", "Bagian Pembangunan dan Sumber Daya Alam", "Bagian Pengadaan Barang dan Jasa", 
                    "Bagian Umum", "Bagian Organisasi", "Bagian Protokol dan Komunikasi Pimpinan", "Bagian Perencanaan dan Keuangan"
                ],
                "Kecamatan": [
                    "Kecamatan Bahar Selatan", "Kecamatan Bahar Utara", "Kecamatan Jambi Luar Kota", "Kecamatan Taman Rajo", 
                    "Kecamatan Kumpeh", "Kecamatan Kumpeh Ulu", "Kecamatan Maro Sebo", "Kecamatan Mestong", "Kecamatan Sekernan", 
                    "Kecamatan Sungai Bahar", "Kecamatan Sungai Gelam"
                ],
                "Rumah Sakit": [
                    "RSUD Ahmad Ripin", "RSUD Sungai Gelam", "RSUD Sungai Bahar"
                ],
                "Puskesmas": [
                    "Puskesmas Penyengat Olak", "Puskesmas Kemingking Dalam", "Puskesmas Puding", "Puskesmas Sungai Bahar IV", 
                    "Puskesmas Simpang Sungai Duren", "Puskesmas Tempino", "Puskesmas Tanjung", "Puskesmas Bahar VII", 
                    "Puskesmas Tantan", "Puskesmas Kasang Pudak", "Puskesmas Sengeti", "Puskesmas Pir II Bajubang", 
                    "Puskesmas Pondok Meja", "Puskesmas Markanding", "Puskesmas Tangkit", "Puskesmas Talang Bukit", 
                    "Puskesmas Sekernan Ilir", "Puskesmas Jambi Kecil", "Puskesmas Muara Kumpeh", "Puskesmas Sungai Bahar I", 
                    "Puskesmas Kebon IX", "Puskesmas Suko Awin", "Puskesmas Petaling Jaya"
                ]
            }

            df_unique = df_filtered.copy()
            df_unique["NIP_Clean"] = df_unique["NIP Admin"].astype(str).str.strip()
            df_unique = df_unique[df_unique["NIP_Clean"] != ""]
            df_unique = df_unique.drop_duplicates(subset=["NIP_Clean"], keep="last")
            
            pns_count = int(df_unique["Nama Admin"].astype(str).str.upper().str.contains("PNS").sum())
            pppk_count = int(df_unique["Nama Admin"].astype(str).str.upper().str.contains("PPPK").sum())
            total_unique = len(df_unique)
            
            total_spt = len(df_filtered)
            sitpp_spt = len(df_filtered[df_filtered["Integrasi"] == "SITPP"])
            other_spt = total_spt - sitpp_spt
            
            # --- Render Statistik KPI ---
            st.write("#### 📊 Statistik Kepegawaian (NIP Unik)")
            c1, c2, c3 = st.columns(3)
            with c1:
                st.metric("Total PNS", pns_count)
            with c2:
                st.metric("Total PPPK", pppk_count)
            with c3:
                st.metric("Total Admin Pengirim", total_unique)
                
            st.write("#### 📝 Statistik Pengiriman Dokumen")
            c4, c5, c6 = st.columns(3)
            with c4:
                st.metric("Total SPT Masuk", total_spt)
            with c5:
                st.metric("Integrasi SiTPP", sitpp_spt)
            with c6:
                st.metric("SPT Lainnya / Tanpa Integrasi", other_spt)
                
            st.write("---")
            
            # --- Render Progress Kategori ---
            st.write("### 🏛️ Progress Unit Terlapor")
            
            all_master_units = set()
            for kat, units in masterUnit.items():
                for u in units:
                    all_master_units.add(u)
                    
            reported_units = set(df_unique["Unit Kerja"].astype(str).str.strip().unique())
            unit_sudah_input = reported_units & all_master_units
            unit_anomali = reported_units - all_master_units - {""}
                    
            cols_prog = st.columns(len(masterUnit) + (1 if len(unit_anomali) > 0 else 0))
            idx_col = 0
            for kat, units in masterUnit.items():
                total_units = len(units)
                sudah_count = sum(1 for u in units if u in unit_sudah_input)
                persen = int((sudah_count / total_units) * 100) if total_units > 0 else 0
                with cols_prog[idx_col]:
                    st.metric(label=kat, value=f"{sudah_count} / {total_units}", delta=f"{persen}%")
                    st.progress(persen / 100.0)
                idx_col += 1
                
            if len(unit_anomali) > 0:
                with cols_prog[idx_col]:
                    st.metric(label="Anomali (Salah Ketik)", value=f"{len(unit_anomali)}", delta="Perlu Cek", delta_color="inverse")
                    st.progress(1.0)
                    
            st.write("---")
            
            # --- Render Status Kehadiran Data (Grid Box) ---
            st.write("### 🚩 Status Kehadiran Data")
            st.caption("✅ Sudah Input | ❌ Belum Ada Input | ⚠️ Nama Unit Tidak Sesuai Master")
            
            st.markdown("""
                <style>
                .unit-grid {
                    display: flex;
                    flex-wrap: wrap;
                    gap: 6px;
                    margin-bottom: 15px;
                }
                .unit-card {
                    font-size: 0.75rem;
                    padding: 5px 10px;
                    border-radius: 6px;
                    font-weight: bold;
                    display: inline-block;
                }
                .sudah {
                    background-color: #d1e7dd;
                    color: #0f5132;
                    border: 1px solid #badbcc;
                }
                .belum {
                    background-color: #f8d7da;
                    color: #842029;
                    border: 1px solid #f5c2c7;
                    opacity: 0.7;
                }
                .anomali {
                    background-color: #fff3cd;
                    color: #856404;
                    border: 1px solid #ffeeba;
                }
                .category-section {
                    background-color: #ffffff;
                    padding: 12px 18px;
                    border-radius: 10px;
                    margin-bottom: 15px;
                    border: 1px solid #e9ecef;
                    box-shadow: 0 2px 4px rgba(0,0,0,0.01);
                }
                .category-title {
                    font-weight: bold;
                    font-size: 0.95em;
                    margin-bottom: 8px;
                    color: #333333;
                    border-left: 4px solid #0d6efd;
                    padding-left: 8px;
                    text-transform: uppercase;
                }
                </style>
            """, unsafe_allow_html=True)
            
            for kat, units in masterUnit.items():
                card_html = []
                for u in units:
                    is_done = u in unit_sudah_input
                    class_name = "sudah" if is_done else "belum"
                    icon = "✅" if is_done else "❌"
                    card_html.append(f'<span class="unit-card {class_name}">{icon} {u}</span>')
                    
                st.markdown(f"""
                    <div class="category-section">
                        <div class="category-title">{kat}</div>
                        <div class="unit-grid">
                            {"".join(card_html)}
                        </div>
                    </div>
                """, unsafe_allow_html=True)
                
            if len(unit_anomali) > 0:
                card_html = []
                for u in unit_anomali:
                    card_html.append(f'<span class="unit-card anomali">⚠️ {u}</span>')
                    
                st.markdown(f"""
                    <div class="category-section" style="border: 1px solid #ffeeba; background-color: #fffdf5;">
                        <div class="category-title" style="color: #856404; border-left-color: #ffc107;">Anomali (Perlu Cek)</div>
                        <div class="unit-grid">
                            {"".join(card_html)}
                        </div>
                    </div>
                """, unsafe_allow_html=True)
                
            st.write("---")
            
            if total_spt > 0:
                col_chart1, col_chart2 = st.columns(2)
                
                opd_counts = df_filtered["Unit Kerja"].value_counts().reset_index()
                opd_counts.columns = ["Unit Kerja", "Jumlah"]
                
                keg_counts = df_filtered["Perihal"].value_counts().reset_index()
                keg_counts.columns = ["Kegiatan", "Jumlah"]
                
                with col_chart1:
                    st.subheader("Distribusi SPT per Unit Kerja / OPD")
                    st.bar_chart(opd_counts.set_index("Unit Kerja"))
                with col_chart2:
                    st.subheader("Distribusi SPT berdasarkan Kegiatan")
                    st.bar_chart(keg_counts.set_index("Kegiatan"))
                    
            st.write("---")
            
            st.subheader("Detail Data Submisi SPT")
            search_query = st.text_input("Cari berdasarkan Nama, NIP, OPD, atau Email:")
            
            if search_query:
                q = search_query.lower()
                df_display = df_filtered[
                    df_filtered["Nama Admin"].astype(str).str.lower().str.contains(q) |
                    df_filtered["NIP Admin"].astype(str).str.lower().str.contains(q) |
                    df_filtered["Unit Kerja"].astype(str).str.lower().str.contains(q) |
                    df_filtered["Email"].astype(str).str.lower().str.contains(q)
                ]
            else:
                df_display = df_filtered
                
            preview_cols = [c for c in df_display.columns if c not in ["Base64 TTD", "ttd_b64", "ttd", "TTD"]]
            
            st.dataframe(df_display[preview_cols], width="stretch")
            
            csv = df_display.to_csv(index=False).encode('utf-8')
            st.download_button(
                label="📥 Unduh Data Rekap (CSV)",
                data=csv,
                file_name=f"Rekap_SPT_{selected_year}.csv",
                mime="text/csv",
                width="stretch"
            )
            
    with tab2:
        st.subheader("Manajemen Daftar Kegiatan / Perihal")
        kegiatan_list = get_kegiatan_list()
        
        with st.form("tambah_kegiatan_form"):
            st.write("**Tambah Kegiatan Baru**")
            nama_baru = st.text_input("Nama Kegiatan (Contoh: Rekon TPP)")
            tipe_deadline = st.radio("Tipe Batas Tanggal:", ["Tanpa Batas", "Pilih Tanggal Kalender"], horizontal=True)
            if tipe_deadline == "Pilih Tanggal Kalender":
                tgl_deadline = st.date_input("Pilih Tanggal:", value=datetime.date.today())
                deadline_baru = format_indo_date(tgl_deadline)
            else:
                deadline_baru = "Tanpa Batas"
            integrasi_baru = st.selectbox("Target Integrasi:", ["None", "SITPP", "Lainnya"])
            kustom_integrasi = st.text_input("Nama Integrasi Kustom (Jika memilih Lainnya):")
            
            submit_keg = st.form_submit_button("Tambah Kegiatan", type="primary")
            if submit_keg:
                if not nama_baru.strip():
                    st.error("Nama kegiatan tidak boleh kosong!")
                else:
                    final_integrasi = kustom_integrasi.strip() if integrasi_baru == "Lainnya" else integrasi_baru
                    new_item = {
                        "nama": nama_baru.strip(),
                        "integrasi": final_integrasi,
                        "status": "Aktif",
                        "deadline": deadline_baru.strip()
                    }
                    kegiatan_list.append(new_item)
                    if save_kegiatan_list(kegiatan_list):
                        st.success(f"Kegiatan '{nama_baru}' berhasil ditambahkan!")
                        st.rerun()
                        
        st.write("---")
        st.write("**Daftar Kegiatan Aktif**")
        
        if not kegiatan_list:
            st.info("Belum ada kegiatan terdaftar.")
        else:
            for idx, k in enumerate(kegiatan_list):
                col_name, col_int, col_deadline, col_action = st.columns([3, 2, 3, 2])
                with col_name:
                    st.write(f"**{k['nama']}**")
                with col_int:
                    st.code(k["integrasi"])
                with col_deadline:
                    st.write(f"Batas: {k.get('deadline', 'Tanpa Batas')}")
                with col_action:
                    if st.button("Hapus", key=f"del_{idx}", type="secondary"):
                        kegiatan_list.pop(idx)
                        if save_kegiatan_list(kegiatan_list):
                            st.success("Kegiatan berhasil dihapus!")
                            st.rerun()

# --- 6. NAVIGASI SIDEBAR ---
st.sidebar.title("Navigasi")
menu = st.sidebar.radio("Pilih Halaman:", ["📝 Kirim SPT", "🔒 Halaman Admin"])

if menu == "📝 Kirim SPT":
    show_operator_form()
else:
    show_admin_page()

# --- 7. FOOTER ---
st.write("")
st.write("---")
st.markdown(
    """
    <div style='text-align: center; color: #808495; font-size: 0.9em;'>
        Made with Love ❤️ oleh <br>
        <strong>Tim Bagian Organisasi Setda Kab. Muaro Jambi #SlavaUkraini</strong>
    </div>
    """, 
    unsafe_allow_html=True
)

